import gc
import hashlib
import json
import os
import queue
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, Iterable

_CWD = Path.cwd()
_HF_HOME = (_CWD / "models" / ".hf_home").resolve()
_HF_HUB_CACHE = (_CWD / "models" / ".hf_hub_cache").resolve()
_HF_MODULES_CACHE = (_CWD / "models" / ".hf_modules").resolve()
os.environ.setdefault("HF_HOME", str(_HF_HOME))
os.environ.setdefault("HF_HUB_CACHE", str(_HF_HUB_CACHE))
os.environ.setdefault("HF_MODULES_CACHE", str(_HF_MODULES_CACHE))
os.environ.setdefault("PYTORCH_CUDA_ALLOC_CONF", "")

import numpy as np  

import torch  

from huggingface_hub import snapshot_download  


for _tk_env in ("TCL_LIBRARY", "TK_LIBRARY", "TCLLIBPATH"):
    os.environ.pop(_tk_env, None)

try:
    import tkinter as tk  

    from tkinter import ttk  

    from tkinter.scrolledtext import ScrolledText  

except ImportError as exc:  

    tk = None  

    ttk = None  

    ScrolledText = object  

    _TK_IMPORT_ERROR = exc
else:
    _TK_IMPORT_ERROR = None

try:
    import faiss  

except ImportError as exc:  

    raise SystemExit(
        "Missing dependency: faiss. Install faiss-cpu or faiss-gpu."
    ) from exc

try:
    from transformers import AutoModel, AutoTokenizer
except ImportError as exc:  

    raise SystemExit(
        "Missing dependency: transformers. Install: pip install transformers accelerate sentencepiece"
    ) from exc

try:
    from llama_cpp import Llama
except ImportError as exc:
    raise SystemExit(
        "Missing dependency: llama-cpp-python. Install the cu124 wheel from https://github.com/abetlen/llama-cpp-python/releases/tag/v0.3.16-cu124"
    ) from exc

if not sys.platform.startswith("linux"):
    raise SystemExit("This build supports Ubuntu/Linux only.")


def _wsl_copy_to_windows(text: str) -> None:
    """Pipe text to Windows clipboard via clip.exe (WSL2 → Windows clipboard bridge)."""
    try:
        subprocess.run(
            ["/mnt/c/Windows/System32/clip.exe"],
            input=text.encode("utf-16le"),
            check=False,
            timeout=2,
        )
    except Exception:
        pass




def _tk_install_hint() -> str:
    return "On Ubuntu install: sudo apt install python3-tk"




DATA_DIR = Path("data")
MODELS_DIR = Path("models")
EMBED_MODEL_NAMES = (
    "BAAI/bge-m3",
    "nomic-ai/nomic-embed-text-v2-moe",
)
NOMIC_TEXT_EMBED_MODEL_NAME = "nomic-ai/nomic-embed-text-v2-moe"
NOMIC_TEXT_EMBED_SUPPORT_REPO = "nomic-ai/nomic-bert-2048"
NOMIC_TEXT_EMBED_REQUIRED_FILES = (
    "config.json",
    "model.safetensors",
    "tokenizer.json",
)
NOMIC_TEXT_EMBED_SUPPORT_FILES = (
    "configuration_hf_nomic_bert.py",
    "modeling_hf_nomic_bert.py",
)
LLM_MODEL_NAME = "bartowski/nvidia_Llama-3.1-Nemotron-Nano-8B-v1-GGUF"
LLM_GGUF_FILE = "nvidia_Llama-3.1-Nemotron-Nano-8B-v1-bf16.gguf"
RAG_CACHE_DIRNAME = ".rag_cache_english_plus_plus"
RAG_INDEX_CACHE_VERSION = 1

CHUNK_SIZE = 1200
CHUNK_OVERLAP = 180
EMBED_MAX_LENGTH = 1024
TEXT_MAX_BYTES = 4_000_000
LLM_MAX_MODEL_LEN = int(os.getenv("RAG_MAX_MODEL_LEN", "4096"))

DOC_PREFIX = "search_document: "
QUERY_PREFIX = "search_query: "
SESSION_LOG_FILENAME = "RAGLLM_English_Plus_Plus_Session.txt"

COLOR_BG = "#17191C"
COLOR_PANEL = "#1D2024"
COLOR_BORDER = "#2A2D32"
COLOR_FG = "#FFFFFF"
COLOR_MUTED = "#A7ABB2"
COLOR_ACTIVE = "#252930"
COLOR_SELECT = "#3A404A"

ALLOWED_SUFFIXES = {
    ".txt",
    ".md",
    ".py",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".ini",
    ".cfg",
    ".csv",
    ".log",
    ".rst",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".hpp",
    ".go",
    ".rs",
    ".sh",
    ".ps1",
    ".bat",
    ".doc",
    ".docx",
}


@dataclass(slots=True)
class Chunk:
    path: str
    text: str


@dataclass(slots=True)
class EmbedModelRuntime:
    name: str
    model_dir: Path
    tokenizer: object
    model: object


def _safe_model_dir_name(model_name: str) -> str:
    return model_name.replace("/", "__").replace("\\", "__").replace(":", "_")


def _model_local_dir(model_name: str, family: str) -> Path:
    return MODELS_DIR / family / _safe_model_dir_name(model_name)


def _ensure_local_model(
    model_name: str,
    family: str,
    status_cb: Callable[[str], None],
    *,
    local_only_if_present: bool = False,
    required_files: tuple[str, ...] | None = None,
    allow_patterns: list[str] | None = None,
) -> Path:
    local_dir = _model_local_dir(model_name, family)
    cache_dir = MODELS_DIR / ".hf_hub_cache"
    local_dir.mkdir(parents=True, exist_ok=True)
    cache_dir.mkdir(parents=True, exist_ok=True)
    if local_only_if_present and _dir_has_files(local_dir, required_files):
        status_cb(f"Using local {model_name} from {local_dir}")
        return local_dir
    status_cb(f"Syncing {model_name} to {local_dir}")
    snapshot_download(
        repo_id=model_name,
        local_dir=local_dir,
        cache_dir=cache_dir,
        local_files_only=False,
        allow_patterns=allow_patterns,
    )
    return local_dir


def _normalize_rows(vectors: np.ndarray) -> np.ndarray:
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    norms = np.clip(norms, a_min=1e-12, a_max=None)
    return (vectors / norms).astype(np.float32, copy=False)


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(len(text), start + chunk_size)
        raw = text[start:end]
        if end < len(text):
            split_pos = raw.rfind("\n")
            if split_pos > chunk_size // 3:
                end = start + split_pos
                raw = text[start:end]
        raw = raw.strip()
        if raw:
            chunks.append(raw)
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


def _iter_text_files(root_dir: Path) -> Iterable[Path]:
    if not root_dir.exists():
        return []
    files = []
    for path in root_dir.rglob("*"):
        if not path.is_file():
            continue
        if any(part.startswith(".rag_cache") for part in path.parts):
            continue
        suffix = path.suffix.lower()
        if suffix and suffix not in ALLOWED_SUFFIXES:
            continue
        files.append(path)
    files.sort()
    return files


def _read_docx(path: Path) -> str | None:
    try:
        from docx import Document  

    except ImportError:
        return None

    try:
        document = Document(str(path))
    except Exception:
        return None

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))

    content = "\n".join(lines).strip()
    return content or None


def _run_capture(args: list[str], timeout_seconds: int = 90) -> str | None:
    try:
        process = subprocess.run(
            args,
            check=False,
            capture_output=True,
            timeout=timeout_seconds,
        )
    except Exception:
        return None

    if process.returncode != 0:
        return None

    output = process.stdout.decode("utf-8", errors="ignore").strip()
    return output or None


def _read_doc_with_libreoffice(path: Path) -> str | None:
    for tool in ("soffice", "libreoffice"):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            try:
                process = subprocess.run(
                    [
                        tool,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        str(temp_path),
                        str(path),
                    ],
                    check=False,
                    capture_output=True,
                    timeout=120,
                )
            except Exception:
                continue

            if process.returncode != 0:
                continue

            converted_files = sorted(temp_path.glob("*.txt"))
            if not converted_files:
                continue

            try:
                text = converted_files[0].read_text(encoding="utf-8", errors="ignore").strip()
            except OSError:
                continue

            if text:
                return text
    return None


def _read_doc_with_cli(path: Path) -> str | None:
    for tool in ("antiword", "catdoc"):
        text = _run_capture([tool, str(path)], timeout_seconds=90)
        if text:
            return text
    return None


def _read_doc(path: Path) -> str | None:
    for reader in (_read_doc_with_libreoffice, _read_doc_with_cli):
        text = reader(path)
        if text:
            return text.strip()
    return None


def _read_office_document(path: Path) -> str | None:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".doc":
        return _read_doc(path)
    return None


def _read_text(path: Path) -> str | None:
    try:
        stat = path.stat()
    except OSError:
        return None
    if stat.st_size == 0 or stat.st_size > TEXT_MAX_BYTES:
        return None
    if path.suffix.lower() in {".doc", ".docx"}:
        return _read_office_document(path)
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if b"\x00" in raw:
        return None
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return None


def _compute_data_signature(root_dir: Path) -> tuple[str, int]:
    hasher = hashlib.sha256()
    source_count = 0
    for path in _iter_text_files(root_dir):
        try:
            stat = path.stat()
        except OSError:
            continue
        try:
            rel_path = path.relative_to(root_dir).as_posix()
        except ValueError:
            rel_path = path.as_posix()
        hasher.update(rel_path.encode("utf-8", errors="ignore"))
        hasher.update(b"\0")
        hasher.update(str(stat.st_size).encode("ascii"))
        hasher.update(b"\0")
        hasher.update(str(stat.st_mtime_ns).encode("ascii"))
        hasher.update(b"\n")
        source_count += 1
    return hasher.hexdigest(), source_count


def _dir_has_files(path: Path, required_files: tuple[str, ...] | None = None) -> bool:
    if not path.exists() or not path.is_dir():
        return False
    if not required_files:
        return any(path.iterdir())
    return all((path / filename).exists() for filename in required_files)


def _offline_mode_enabled() -> bool:
    flags = (
        os.getenv("HF_HUB_OFFLINE", ""),
        os.getenv("TRANSFORMERS_OFFLINE", ""),
    )
    return any(flag.strip().lower() in {"1", "true", "yes", "on"} for flag in flags)


def _copy_support_files(src_dir: Path, dst_dir: Path, filenames: tuple[str, ...]) -> None:
    dst_dir.mkdir(parents=True, exist_ok=True)
    for filename in filenames:
        src = src_dir / filename
        if not src.exists():
            raise RuntimeError(f"Missing support file in {src_dir}: {filename}")
        shutil.copy2(src, dst_dir / filename)


def _copy_support_files_from_hf_modules(
    support_dir: Path,
    filenames: tuple[str, ...],
) -> bool:
    base = _HF_MODULES_CACHE / "transformers_modules" / "nomic_hyphen_ai" / "nomic_hyphen_bert_hyphen_2048"
    if not base.exists():
        return False
    for candidate in sorted(base.glob("*"), key=lambda p: p.stat().st_mtime_ns, reverse=True):
        if not candidate.is_dir():
            continue
        if _dir_has_files(candidate, filenames):
            _copy_support_files(candidate, support_dir, filenames)
            return True
    return False


def _ensure_nomic_text_embed_local_support(model_dir: Path, status_cb: Callable[[str], None]) -> None:
    support_dir = _model_local_dir(NOMIC_TEXT_EMBED_SUPPORT_REPO, "embeddings")
    support_cache = MODELS_DIR / ".hf_hub_cache"
    support_dir.mkdir(parents=True, exist_ok=True)
    support_cache.mkdir(parents=True, exist_ok=True)

    if not _dir_has_files(model_dir, NOMIC_TEXT_EMBED_SUPPORT_FILES):
        if not _dir_has_files(support_dir, NOMIC_TEXT_EMBED_SUPPORT_FILES):
            if _copy_support_files_from_hf_modules(support_dir, NOMIC_TEXT_EMBED_SUPPORT_FILES):
                status_cb(f"Using local HF modules support code for {NOMIC_TEXT_EMBED_MODEL_NAME}.")
            elif _offline_mode_enabled():
                raise RuntimeError(
                    f"Missing local support code for {NOMIC_TEXT_EMBED_MODEL_NAME}. "
                    f"Run once online to sync {NOMIC_TEXT_EMBED_SUPPORT_REPO}."
                )
            else:
                status_cb(f"Syncing {NOMIC_TEXT_EMBED_SUPPORT_REPO} support code to {support_dir}")
                snapshot_download(
                    repo_id=NOMIC_TEXT_EMBED_SUPPORT_REPO,
                    local_dir=support_dir,
                    cache_dir=support_cache,
                    allow_patterns=list(NOMIC_TEXT_EMBED_SUPPORT_FILES),
                    local_files_only=False,
                )
        _copy_support_files(support_dir, model_dir, NOMIC_TEXT_EMBED_SUPPORT_FILES)

    config_path = model_dir / "config.json"
    if not config_path.exists():
        raise RuntimeError(f"Missing config for {NOMIC_TEXT_EMBED_MODEL_NAME}: {config_path}")
    try:
        config = json.loads(config_path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise RuntimeError(f"Invalid config JSON for {NOMIC_TEXT_EMBED_MODEL_NAME}: {config_path}") from exc

    auto_map = config.get("auto_map")
    if not isinstance(auto_map, dict):
        return

    changed = False
    patched_auto_map: dict[str, object] = {}
    for key, value in auto_map.items():
        if isinstance(value, str) and "--" in value:
            patched_auto_map[key] = value.split("--", 1)[1]
            changed = True
        else:
            patched_auto_map[key] = value
    if changed:
        config["auto_map"] = patched_auto_map
        config_path.write_text(json.dumps(config, indent=2), encoding="utf-8")
        status_cb(f"Patched {NOMIC_TEXT_EMBED_MODEL_NAME} config for local-only remote code.")


class Embedder:
    def __init__(self, status_cb: Callable[[str], None]) -> None:
        self.status_cb = status_cb
        self.device = self._select_device()
        self.batch_size = int(os.getenv("RAG_EMBED_BATCH", "4"))
        self.runtimes: list[EmbedModelRuntime] = []
        dtype = torch.float16 if self.device.type == "cuda" else torch.float32
        for model_name in EMBED_MODEL_NAMES:
            if model_name == NOMIC_TEXT_EMBED_MODEL_NAME:
                model_dir = _ensure_local_model(
                    model_name,
                    "embeddings",
                    self.status_cb,
                    local_only_if_present=True,
                    required_files=NOMIC_TEXT_EMBED_REQUIRED_FILES,
                )
                _ensure_nomic_text_embed_local_support(model_dir, self.status_cb)
            else:
                model_dir = _ensure_local_model(
                    model_name,
                    "embeddings",
                    self.status_cb,
                    local_only_if_present=True,
                    required_files=("config.json", "tokenizer.json"),
                )
            self.status_cb(f"Loading embedding model on {self.device.type}: {model_dir}")
            tokenizer = AutoTokenizer.from_pretrained(
                str(model_dir),
                trust_remote_code=True,
                local_files_only=True,
            )
            model = AutoModel.from_pretrained(
                str(model_dir),
                trust_remote_code=True,
                dtype=dtype,
                low_cpu_mem_usage=True,
                local_files_only=True,
            )
            model.to(self.device)
            model.eval()
            self.runtimes.append(
                EmbedModelRuntime(
                    name=model_name,
                    model_dir=model_dir,
                    tokenizer=tokenizer,
                    model=model,
                )
            )
        if not self.runtimes:
            raise RuntimeError("No embedding models were loaded.")

    @staticmethod
    def _select_device() -> torch.device:
        override = os.getenv("RAG_EMBED_DEVICE", "").strip().lower()
        if override == "cpu":
            raise RuntimeError("CPU embedding is disabled in Linux CUDA mode.")
        if override and override != "cuda":
            raise RuntimeError("RAG_EMBED_DEVICE must be unset or set to 'cuda'.")
        if not torch.cuda.is_available():
            raise RuntimeError("CUDA GPU is required. Install NVIDIA driver/CUDA and use a CUDA-enabled environment.")
        return torch.device("cuda")

    def _pool_embeddings(self, model_out: object, attention_mask: torch.Tensor) -> torch.Tensor:
        if torch.is_tensor(model_out):
            return model_out
        if isinstance(model_out, dict):
            if "sentence_embedding" in model_out:
                return model_out["sentence_embedding"]
            if "embeddings" in model_out:
                return model_out["embeddings"]
            last_hidden = model_out.get("last_hidden_state")
        else:
            last_hidden = getattr(model_out, "last_hidden_state", None)
            if last_hidden is None and isinstance(model_out, tuple) and model_out:
                last_hidden = model_out[0]

        if last_hidden is None:
            raise RuntimeError("Embedding model output does not contain usable embeddings.")

        mask = attention_mask.unsqueeze(-1).to(last_hidden.dtype)
        masked = last_hidden * mask
        summed = masked.sum(dim=1)
        denom = torch.clamp(mask.sum(dim=1), min=1e-9)
        return summed / denom

    def _encode_with_runtime(self, runtime: EmbedModelRuntime, texts: list[str]) -> np.ndarray:
        encoded = runtime.tokenizer(
            texts,
            padding=True,
            truncation=True,
            max_length=EMBED_MAX_LENGTH,
            return_tensors="pt",
        )
        encoded = {k: v.to(self.device) for k, v in encoded.items()}
        with torch.inference_mode():
            out = runtime.model(**encoded)
            vectors = self._pool_embeddings(out, encoded["attention_mask"])
        return torch.nn.functional.normalize(vectors.float(), dim=1).detach().cpu().numpy()

    def _encode(self, texts: list[str]) -> np.ndarray:
        per_model_vectors: list[np.ndarray] = []
        for runtime in self.runtimes:
            per_model_vectors.append(self._encode_with_runtime(runtime, texts))
        if not per_model_vectors:
            raise RuntimeError("No embedding vectors generated.")
        if len(per_model_vectors) == 1:
            return per_model_vectors[0]
        combined = np.concatenate(per_model_vectors, axis=1)
        return _normalize_rows(combined)

    def encode_documents(self, texts: list[str]) -> np.ndarray:
        prefixed = [DOC_PREFIX + t for t in texts]
        all_vectors = []
        for start in range(0, len(prefixed), self.batch_size):
            batch = prefixed[start : start + self.batch_size]
            all_vectors.append(self._encode(batch))
        return np.vstack(all_vectors).astype(np.float32, copy=False)

    def encode_query(self, query: str) -> np.ndarray:
        return self._encode([QUERY_PREFIX + query])


class RAGEngine:
    def __init__(self, status_cb: Callable[[str], None]) -> None:
        self.status_cb = status_cb
        self.lock = threading.Lock()
        self.ready = False

        self.embedder: Embedder | None = None
        self.llm_engine = None

        self.chunks: list[Chunk] = []
        self.index_cpu: faiss.Index | None = None
        self.index_runtime: faiss.Index | None = None
        self.faiss_gpu_resources = None

    def _configure_cuda(self) -> None:
        if not torch.cuda.is_available():
            raise RuntimeError("CUDA GPU is required. CPU-only execution is disabled.")
        torch.set_float32_matmul_precision("high")
        torch.backends.cuda.matmul.allow_tf32 = True
        torch.backends.cudnn.allow_tf32 = True
        torch.backends.cudnn.benchmark = True
        cudnn_version = torch.backends.cudnn.version()
        total_gib = torch.cuda.get_device_properties(0).total_memory / (1024**3)
        self.status_cb(f"CUDA ready: {torch.cuda.get_device_name(0)} | VRAM {total_gib:.1f} GiB | cuDNN {cudnn_version}")

    def _load_documents(self) -> list[Chunk]:
        self.status_cb(f"Scanning documents in {DATA_DIR}")
        chunks: list[Chunk] = []
        for path in _iter_text_files(DATA_DIR):
            text = _read_text(path)
            if not text:
                continue
            for part in _chunk_text(text):
                chunks.append(Chunk(path=str(path), text=part))
        if not chunks:
            raise RuntimeError("No text chunks found under data/.")
        self.status_cb(f"Prepared {len(chunks)} chunks for embedding.")
        return chunks

    def _build_faiss_index(self) -> None:
        if self.embedder is None:
            raise RuntimeError("Embedder is not loaded.")

        data_signature, source_count = _compute_data_signature(DATA_DIR)
        if self._try_load_cached_index(data_signature, source_count):
            return

        self.chunks = self._load_documents()
        self.status_cb("Building FAISS index...")
        index_cpu: faiss.Index | None = None

        build_batch_size = int(os.getenv("RAG_INDEX_BATCH", "64"))
        for start in range(0, len(self.chunks), build_batch_size):
            group = self.chunks[start : start + build_batch_size]
            vectors = self.embedder.encode_documents([c.text for c in group])
            if index_cpu is None:
                index_cpu = faiss.IndexFlatIP(vectors.shape[1])
            index_cpu.add(vectors)
            self.status_cb(f"Indexed {min(start + build_batch_size, len(self.chunks))}/{len(self.chunks)} chunks")

        if index_cpu is None:
            raise RuntimeError("Failed to build FAISS index.")
        self._set_runtime_index(index_cpu)
        self._save_cached_index(data_signature, source_count)

    @staticmethod
    def _cache_dir() -> Path:
        return DATA_DIR / RAG_CACHE_DIRNAME

    @classmethod
    def _cache_paths(cls) -> tuple[Path, Path, Path]:
        cache_dir = cls._cache_dir()
        return cache_dir / "meta.json", cache_dir / "chunks.jsonl", cache_dir / "index.faiss"

    @staticmethod
    def _cache_settings_payload() -> dict[str, object]:
        return {
            "cache_version": RAG_INDEX_CACHE_VERSION,
            "embed_models": list(EMBED_MODEL_NAMES),
            "embed_merge": "concat_l2norm",
            "chunk_size": CHUNK_SIZE,
            "chunk_overlap": CHUNK_OVERLAP,
            "embed_max_length": EMBED_MAX_LENGTH,
            "doc_prefix": DOC_PREFIX,
            "query_prefix": QUERY_PREFIX,
        }

    def _set_runtime_index(self, index_cpu: faiss.Index) -> None:
        self.index_cpu = index_cpu
        self.index_runtime = index_cpu
        if not torch.cuda.is_available():
            self.status_cb("CUDA not available for FAISS. Using CPU FAISS index.")
            return
        if not hasattr(faiss, "StandardGpuResources"):
            self.status_cb("FAISS GPU backend not available. Using CPU FAISS index.")
            return
        try:
            self.faiss_gpu_resources = faiss.StandardGpuResources()
            _use_cuvs = False
            if hasattr(faiss, "GpuClonerOptions"):
                options = faiss.GpuClonerOptions()
                options.useFloat16 = True
                if hasattr(options, "use_cuvs"):
                    options.use_cuvs = True
                    _use_cuvs = True
                try:
                    self.index_runtime = faiss.index_cpu_to_gpu(self.faiss_gpu_resources, 0, index_cpu, options)
                except TypeError:
                    self.index_runtime = faiss.index_cpu_to_gpu(self.faiss_gpu_resources, 0, index_cpu)
            else:
                self.index_runtime = faiss.index_cpu_to_gpu(self.faiss_gpu_resources, 0, index_cpu)
            self.status_cb(f"FAISS {'cuVS' if _use_cuvs else 'GPU'} index enabled.")
        except Exception as exc:
            self.index_runtime = index_cpu
            self.status_cb(f"FAISS GPU init failed ({exc}). Falling back to CPU FAISS index.")

    def _try_load_cached_index(self, data_signature: str, source_count: int) -> bool:
        meta_path, chunks_path, index_path = self._cache_paths()
        if not (meta_path.exists() and chunks_path.exists() and index_path.exists()):
            return False

        try:
            meta = json.loads(meta_path.read_text(encoding="utf-8"))
        except Exception:
            return False

        settings = self._cache_settings_payload()
        for key, value in settings.items():
            if meta.get(key) != value:
                return False
        if meta.get("data_signature") != data_signature:
            return False
        if int(meta.get("source_file_count", -1)) != source_count:
            return False

        try:
            chunks: list[Chunk] = []
            with chunks_path.open("r", encoding="utf-8") as handle:
                for line in handle:
                    line = line.strip()
                    if not line:
                        continue
                    item = json.loads(line)
                    chunks.append(Chunk(path=str(item["path"]), text=str(item["text"])))
            index_cpu = faiss.read_index(str(index_path))
            if int(index_cpu.ntotal) != len(chunks):
                raise RuntimeError("chunk/index size mismatch")
        except Exception as exc:
            self.status_cb(f"Cached index invalid ({exc}). Rebuilding.")
            return False

        self.chunks = chunks
        self._set_runtime_index(index_cpu)
        self.status_cb(f"Loaded cached FAISS index ({len(chunks)} chunks).")
        return True

    def _save_cached_index(self, data_signature: str, source_count: int) -> None:
        if self.index_cpu is None:
            return

        cache_dir = self._cache_dir()
        cache_dir.mkdir(parents=True, exist_ok=True)
        meta_path, chunks_path, index_path = self._cache_paths()
        tmp_meta = meta_path.with_suffix(meta_path.suffix + ".tmp")
        tmp_chunks = chunks_path.with_suffix(chunks_path.suffix + ".tmp")
        tmp_index = index_path.with_suffix(index_path.suffix + ".tmp")

        payload = self._cache_settings_payload()
        payload.update(
            {
                "data_signature": data_signature,
                "source_file_count": source_count,
                "chunk_count": len(self.chunks),
                "index_ntotal": int(self.index_cpu.ntotal),
            }
        )

        try:
            with tmp_chunks.open("w", encoding="utf-8") as handle:
                for chunk in self.chunks:
                    handle.write(json.dumps({"path": chunk.path, "text": chunk.text}))
                    handle.write("\n")
            faiss.write_index(self.index_cpu, str(tmp_index))
            tmp_meta.write_text(json.dumps(payload, indent=2), encoding="utf-8")
            tmp_chunks.replace(chunks_path)
            tmp_index.replace(index_path)
            tmp_meta.replace(meta_path)
            self.status_cb(f"Saved FAISS cache to {cache_dir}")
        except Exception as exc:
            self.status_cb(f"Warning: failed to save FAISS cache ({exc})")
            for temp_path in (tmp_chunks, tmp_index, tmp_meta):
                try:
                    if temp_path.exists():
                        temp_path.unlink()
                except OSError:
                    pass

    def _start_llm_load_monitor(self, model_dir: Path) -> threading.Event:
        total_bytes = 0
        for pattern in ("*.safetensors", "*.bin"):
            files = list(model_dir.glob(pattern))
            if files:
                total_bytes = sum(f.stat().st_size for f in files)
                break
        stop = threading.Event()
        if not total_bytes or not torch.cuda.is_available():
            return stop
        baseline = torch.cuda.memory_allocated(0)
        def _run() -> None:
            while not stop.wait(0.3):
                delta = max(0, torch.cuda.memory_allocated(0) - baseline)
                pct = min(99, int(100 * delta / total_bytes))
                self.status_cb(f"__llm_load_pct__:{pct}")
        threading.Thread(target=_run, daemon=True).start()
        return stop

    def _load_llm(self) -> None:
        model_dir = _ensure_local_model(
            LLM_MODEL_NAME, "llm", self.status_cb,
            allow_patterns=[LLM_GGUF_FILE] if LLM_GGUF_FILE else None,
        )
        self.status_cb(f"Loading LLM: {model_dir}")
        if not torch.cuda.is_available():
            raise RuntimeError("CUDA GPU is required for LLM loading.")
        torch.cuda.empty_cache()
        _stop = self._start_llm_load_monitor(model_dir)
        try:
            self.llm_engine = Llama(
                model_path=str(model_dir / LLM_GGUF_FILE),
                n_gpu_layers=-1,
                n_ctx=LLM_MAX_MODEL_LEN,
                verbose=False,
            )
        finally:
            _stop.set()
            self.status_cb("__llm_load_pct__:100")
        self.status_cb("LLM loaded.")

    def prepare(self) -> None:
        if self.ready:
            return
        with self.lock:
            if self.ready:
                return
            self._configure_cuda()
            self.embedder = Embedder(self.status_cb)
            self._build_faiss_index()
            self._load_llm()
            self.ready = True
            self.status_cb("RAG engine ready.")

    def _retrieve(self, query: str, top_k: int) -> list[tuple[float, Chunk]]:
        if self.embedder is None or self.index_runtime is None:
            raise RuntimeError("Retriever is not ready.")

        q_vec = self.embedder.encode_query(query)
        scores, ids = self.index_runtime.search(q_vec, top_k)
        results: list[tuple[float, Chunk]] = []
        for score, idx in zip(scores[0], ids[0], strict=False):
            if idx < 0:
                continue
            results.append((float(score), self.chunks[idx]))
        return results

    def _build_prompt(self, query: str, contexts: list[tuple[float, Chunk]]) -> str:
        context_lines = []
        for i, (score, chunk) in enumerate(contexts, start=1):
            context_lines.append(f"[{i}] source={chunk.path} score={score:.4f}\n{chunk.text}")
        context_text = "\n\n".join(context_lines)
        return (
            "You are a concise assistant. Answer using the provided context. "
            "If the context is insufficient, say you do not know.\n\n"
            f"Context:\n{context_text}\n\nQuestion:\n{query}\n\nAnswer:"
        )

    def _generate(self, prompt: str, max_new_tokens: int, temperature: float) -> str:
        if self.llm_engine is None:
            raise RuntimeError("LLM is not loaded.")
        do_sample = temperature > 0.0
        result = self.llm_engine.create_chat_completion(
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful assistant. "
                        "Think step-by-step before responding. "
                        "Answer using the provided context. "
                        "If the context is insufficient, say you do not know."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_new_tokens,
            temperature=temperature if do_sample else 0.0,
            top_p=0.9 if do_sample else 1.0,
        )
        return result["choices"][0]["message"]["content"].strip()

    def answer(self, query: str, top_k: int, max_new_tokens: int, temperature: float) -> tuple[str, list[tuple[float, Chunk]]]:
        self.prepare()
        contexts = self._retrieve(query, top_k=top_k)
        prompt = self._build_prompt(query, contexts)
        answer = self._generate(prompt, max_new_tokens=max_new_tokens, temperature=temperature)
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        return answer, contexts


class RAGApp:
    def __init__(self) -> None:
        if _TK_IMPORT_ERROR is not None or tk is None or ttk is None:
            raise RuntimeError(f"Missing tkinter support. {_tk_install_hint()}") from _TK_IMPORT_ERROR
        self.root = tk.Tk()
        self.root.title("RAG LLM (FAISS + CUDA)")
        self.root.geometry("1200x820")
        self.root.configure(bg=COLOR_BG)
        self._apply_dark_theme()

        self.status_var = tk.StringVar(master=self.root, value="Starting...")
        self.indexing_status_var = tk.StringVar(master=self.root, value="Indexing chunks: waiting...")
        self._indexing_started_at: float | None = None
        self.ui_queue: queue.Queue[tuple[str, object]] = queue.Queue()
        self.session_log_path = Path.cwd() / SESSION_LOG_FILENAME

        self.engine = RAGEngine(self._post_status)
        self.query_thread: threading.Thread | None = None
        self.boot_thread: threading.Thread | None = None
        self._start_session_log()

        self._build_ui()
        self.root.after(120, self._drain_queue)
        self._start_bootstrap()

    def _apply_dark_theme(self) -> None:
        style = ttk.Style(self.root)
        try:
            style.theme_use("clam")
        except tk.TclError:
            pass

        style.configure(".", background=COLOR_BG, foreground=COLOR_FG)
        style.configure("TFrame", background=COLOR_BG)
        style.configure("TLabel", background=COLOR_BG, foreground=COLOR_FG)
        style.configure(
            "TEntry",
            foreground=COLOR_FG,
            fieldbackground=COLOR_PANEL,
            background=COLOR_PANEL,
            bordercolor=COLOR_BORDER,
            lightcolor=COLOR_BORDER,
            darkcolor=COLOR_BORDER,
        )
        style.map(
            "TEntry",
            fieldbackground=[("readonly", COLOR_PANEL), ("disabled", COLOR_BG)],
            foreground=[("disabled", COLOR_MUTED)],
        )
        style.configure(
            "TButton",
            foreground=COLOR_FG,
            background=COLOR_PANEL,
            bordercolor=COLOR_BORDER,
            lightcolor=COLOR_BORDER,
            darkcolor=COLOR_BORDER,
            focuscolor=COLOR_BORDER,
        )
        style.map(
            "TButton",
            background=[("active", COLOR_ACTIVE), ("pressed", COLOR_ACTIVE), ("disabled", COLOR_BG)],
            foreground=[("disabled", COLOR_MUTED)],
        )
        style.configure(
            "Vertical.TScrollbar",
            background=COLOR_PANEL,
            troughcolor=COLOR_BG,
            bordercolor=COLOR_BG,
            arrowcolor=COLOR_FG,
        )
        style.configure(
            "Horizontal.TProgressbar",
            troughcolor=COLOR_PANEL,
            background="#4A9EFF",
            bordercolor=COLOR_BORDER,
            lightcolor=COLOR_BORDER,
            darkcolor=COLOR_BORDER,
        )

    def _style_scrolled_text(self, widget: ScrolledText) -> None:
        widget.configure(
            bg=COLOR_PANEL,
            fg=COLOR_FG,
            insertbackground=COLOR_FG,
            selectbackground=COLOR_SELECT,
            selectforeground=COLOR_FG,
            relief="flat",
            borderwidth=1,
            highlightthickness=1,
            highlightbackground=COLOR_BORDER,
            highlightcolor=COLOR_BORDER,
        )
        try:
            widget.vbar.configure(
                bg=COLOR_PANEL,
                activebackground=COLOR_ACTIVE,
                troughcolor=COLOR_BG,
                highlightthickness=0,
                relief="flat",
            )
        except tk.TclError:
            pass

    def _build_ui(self) -> None:
        main = ttk.Frame(self.root, padding=12)
        main.pack(fill="both", expand=True)

        top = ttk.Frame(main)
        top.pack(fill="x")

        ttk.Label(top, text="Question").grid(row=0, column=0, sticky="w")
        self.top_k_var = tk.StringVar(master=self.root, value="7")
        self.max_tokens_var = tk.StringVar(master=self.root, value="4096")
        self.temp_var = tk.StringVar(master=self.root, value="0.7")

        ttk.Label(top, text="Top K").grid(row=0, column=1, padx=(12, 4))
        ttk.Entry(top, textvariable=self.top_k_var, width=6).grid(row=0, column=2)
        ttk.Label(top, text="Max New Tokens").grid(row=0, column=3, padx=(12, 4))
        ttk.Entry(top, textvariable=self.max_tokens_var, width=8).grid(row=0, column=4)
        ttk.Label(top, text="Temperature").grid(row=0, column=5, padx=(12, 4))
        ttk.Entry(top, textvariable=self.temp_var, width=6).grid(row=0, column=6)

        self.ask_btn = ttk.Button(top, text="Ask", command=self._on_ask)
        self.ask_btn.grid(row=0, column=7, padx=(12, 0))

        self.question_box = ScrolledText(main, height=7, wrap="word")
        self._style_scrolled_text(self.question_box)
        self.question_box.pack(fill="x", pady=(8, 8))

        ttk.Label(main, text="Answer").pack(anchor="w")
        self.answer_box = ScrolledText(main, height=16, wrap="word")
        self._style_scrolled_text(self.answer_box)
        self.answer_box.pack(fill="both", expand=True, pady=(4, 8))

        ttk.Label(main, text="Retrieved Context").pack(anchor="w")
        self.context_box = ScrolledText(main, height=12, wrap="word")
        self._style_scrolled_text(self.context_box)
        self.context_box.pack(fill="both", expand=True, pady=(4, 8))

        tk.Entry(
            main, textvariable=self.status_var, state="readonly",
            readonlybackground=COLOR_BG, foreground=COLOR_FG,
            relief="flat", borderwidth=0, highlightthickness=0, insertwidth=0,
        ).pack(fill="x", anchor="w")
        tk.Entry(
            main, textvariable=self.indexing_status_var, state="readonly",
            readonlybackground=COLOR_BG, foreground=COLOR_FG,
            relief="flat", borderwidth=0, highlightthickness=0, insertwidth=0,
        ).pack(fill="x", anchor="w")
        _pbar_frame = ttk.Frame(main)
        _pbar_frame.pack(fill="x", pady=(2, 0))
        self._llm_progress_bar = ttk.Progressbar(_pbar_frame, mode="determinate", maximum=100, value=0)
        self._llm_progress_bar.pack(side="left", fill="x", expand=True)
        self._llm_progress_label = ttk.Label(_pbar_frame, text="  0%", width=5)
        self._llm_progress_label.pack(side="left", padx=(6, 0))

        def _sync_win_clipboard(event: object) -> None:
            try:
                w = self.root.focus_get()
                text = w.selection_get() if w is not None else ""
                if text:
                    _wsl_copy_to_windows(text)
            except Exception:
                pass
        self.root.bind_all("<Control-c>", _sync_win_clipboard, add=True)
        self.root.bind_all("<<Copy>>", _sync_win_clipboard, add=True)

    def _start_bootstrap(self) -> None:
        def worker() -> None:
            try:
                self.engine.prepare()
            except Exception as exc:
                self.ui_queue.put(("status", f"Bootstrap failed: {exc}"))

        self.boot_thread = threading.Thread(target=worker, daemon=True)
        self.boot_thread.start()

    def _post_status(self, message: str) -> None:
        self.ui_queue.put(("status", message))

    def _start_session_log(self) -> None:
        started = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        try:
            with self.session_log_path.open("a", encoding="utf-8") as handle:
                handle.write(f"\n=== Session started {started} ===\n\n")
        except OSError:
            pass

    def _append_session_entry(self, query: str, answer: str) -> None:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        try:
            with self.session_log_path.open("a", encoding="utf-8") as handle:
                handle.write(f"[{timestamp}] User Query\n")
                handle.write(query.strip())
                handle.write("\n\n")
                handle.write(f"[{timestamp}] Assistant Answer\n")
                handle.write(answer.strip())
                handle.write("\n\n")
                handle.write("=" * 72)
                handle.write("\n\n")
        except OSError as exc:
            self.ui_queue.put(("status", f"Warning: failed to write session log ({exc})"))

    @staticmethod
    def _format_elapsed(seconds: float) -> str:
        total_seconds = max(0, int(seconds))
        minutes, secs = divmod(total_seconds, 60)
        return f"{minutes:02d}:{secs:02d}"

    def _update_indexing_status(self, message: str) -> None:
        if message == "Building FAISS index...":
            self._indexing_started_at = time.perf_counter()
            self.indexing_status_var.set("Indexing chunks: 0/? | elapsed 00:00")
            return

        if message.startswith("Loaded cached FAISS index"):
            self._indexing_started_at = None
            self.indexing_status_var.set("Indexing chunks: skipped (cache hit).")
            return

        if not (message.startswith("Indexed ") and message.endswith(" chunks")):
            return

        counts = message[len("Indexed ") : -len(" chunks")]
        if "/" not in counts:
            return

        done_str, total_str = counts.split("/", 1)
        try:
            done = int(done_str)
            total = int(total_str)
        except ValueError:
            return

        if self._indexing_started_at is None:
            self._indexing_started_at = time.perf_counter()
        elapsed = max(0.0, time.perf_counter() - self._indexing_started_at)
        if done >= total:
            self.indexing_status_var.set(
                f"Indexing chunks: {done}/{total} | elapsed {self._format_elapsed(elapsed)} (complete)"
            )
            self._indexing_started_at = None
            return
        self.indexing_status_var.set(
            f"Indexing chunks: {done}/{total} | elapsed {self._format_elapsed(elapsed)}"
        )

    def _drain_queue(self) -> None:
        while True:
            try:
                kind, payload = self.ui_queue.get_nowait()
            except queue.Empty:
                break

            if kind == "status":
                message = str(payload)
                if message.startswith("__llm_load_pct__:"):
                    pct = int(message.split(":", 1)[1])
                    self._llm_progress_bar["value"] = pct
                    self._llm_progress_label.configure(text=f"{pct:3d}%")
                else:
                    self.status_var.set(message)
                    self._update_indexing_status(message)
            elif kind == "answer":
                answer, contexts = payload  

                self.answer_box.delete("1.0", "end")
                self.answer_box.insert("1.0", answer)

                self.context_box.delete("1.0", "end")
                lines = []
                for i, (score, chunk) in enumerate(contexts, start=1):
                    lines.append(f"[{i}] score={score:.4f} source={chunk.path}\n{chunk.text}")
                self.context_box.insert("1.0", "\n\n".join(lines))
                self.ask_btn.configure(state="normal")
            elif kind == "error":
                self.answer_box.delete("1.0", "end")
                self.answer_box.insert("1.0", str(payload))
                self.ask_btn.configure(state="normal")

        self.root.after(120, self._drain_queue)

    def _on_ask(self) -> None:
        if self.query_thread and self.query_thread.is_alive():
            self.status_var.set("Query already running...")
            return

        query = self.question_box.get("1.0", "end").strip()
        if not query:
            self.status_var.set("Enter a question first.")
            return

        try:
            top_k = max(1, int(self.top_k_var.get().strip()))
            max_new_tokens = max(32, int(self.max_tokens_var.get().strip()))
            temperature = max(0.0, float(self.temp_var.get().strip()))
        except ValueError:
            self.status_var.set("Top K, Max New Tokens, and Temperature must be numeric.")
            return

        self.ask_btn.configure(state="disabled")
        self.status_var.set("Running retrieval and generation...")

        def worker() -> None:
            try:
                answer, contexts = self.engine.answer(
                    query=query,
                    top_k=top_k,
                    max_new_tokens=max_new_tokens,
                    temperature=temperature,
                )
                self._append_session_entry(query, answer)
                self.ui_queue.put(("answer", (answer, contexts)))
                self.ui_queue.put(("status", "Completed."))
            except Exception as exc:
                self.ui_queue.put(("error", f"Query failed: {exc}"))
                self.ui_queue.put(("status", "Failed."))

        self.query_thread = threading.Thread(target=worker, daemon=True)
        self.query_thread.start()

    def run(self) -> None:
        self.root.mainloop()


def main() -> None:
    app = RAGApp()
    app.run()


if __name__ == "__main__":
    main()
